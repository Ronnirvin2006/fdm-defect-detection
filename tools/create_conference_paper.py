import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "outputs" / "paper" / "FDM_Defect_Detection_Conference_Paper.md"
OUTPUT = ROOT / "outputs" / "paper" / "FDM_Defect_Detection_Conference_Paper.docx"
CONFUSION = ROOT / "outputs" / "figures" / "confusion_matrix.png"
GRADCAM = ROOT / "outputs" / "figures" / "gradcam_Image_20231128195336980.png"


def set_font(run, size=9, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_inline(paragraph, text, size=9.5):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            set_font(paragraph.add_run(part[1:-1]), size=size, italic=True)
        else:
            set_font(paragraph.add_run(part), size=size)


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)


def set_columns(section, count=2, space_twips="360"):
    sect_pr = section._sectPr
    columns = sect_pr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        sect_pr.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), space_twips)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def add_figure(document, image_path: Path, caption: str):
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(image_path), width=Inches(3.35))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(5)
    set_font(cap.add_run(caption), size=8, italic=True)


def add_markdown_table(document, rows):
    headers = [cell.strip() for cell in rows[0].strip("|").split("|")]
    body = rows[2:]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    if len(headers) == 3:
        widths = [Inches(1.12), Inches(1.72), Inches(0.42)]
    else:
        widths = [Inches(3.25 / len(headers)) for _ in headers]
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(int(width.inches * 1440)))
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(paragraph.add_run(value), size=7.5, bold=True)
    for source_row in body:
        values = [cell.strip() for cell in source_row.strip("|").split("|")]
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].width = widths[index]
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_font(paragraph.add_run(value), size=7.5)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    document = Document()
    configure_page(document.sections[0])
    document.core_properties.title = "An Explainable Real-Time Vision System for Multi-Class Defect Detection in FDM 3D Printing Using EfficientNetB0"
    document.core_properties.author = "Ron Nirvin"
    document.core_properties.subject = "IEEE-style conference paper submission draft"

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Paper Caption" not in document.styles:
        caption_style = document.styles.add_style("Paper Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Times New Roman"
        caption_style.font.size = Pt(8)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    two_columns_started = False
    references_mode = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("| "):
            table_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(lines[index].strip())
                index += 1
            add_markdown_table(document, table_rows)
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(8)
            set_font(paragraph.add_run(stripped[2:]), size=16, bold=True)
        elif stripped == "## I. Introduction" and not two_columns_started:
            section = document.add_section(WD_SECTION.CONTINUOUS)
            configure_page(section)
            set_columns(section, 2)
            two_columns_started = True
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            set_font(paragraph.add_run("I. INTRODUCTION"), size=10, bold=True)
        elif stripped.startswith("## "):
            heading = stripped[3:]
            references_mode = heading == "References"
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.keep_with_next = True
            set_font(paragraph.add_run(heading.upper()), size=10, bold=True)
            if heading == "IV. Experimental Results":
                add_figure(document, CONFUSION, "Fig. 1. Confusion matrix for the 661-image held-out test split.")
        elif stripped.startswith("### "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            set_font(paragraph.add_run(stripped[4:]), size=9, bold=True, italic=True)
            if stripped.startswith("### D. Explainability"):
                add_figure(document, GRADCAM, "Fig. 2. Example Grad-CAM overlay for model explanation.")
        elif stripped.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(5)
            set_font(paragraph.add_run(stripped[2:]), size=8, italic=True)
        elif re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            add_inline(paragraph, re.sub(r"^\d+\. ", "", stripped), size=9.5)
        else:
            paragraph = document.add_paragraph()
            if len(document.paragraphs) <= 4:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if references_mode or re.match(r"^\[\d+\]", stripped):
                paragraph.paragraph_format.left_indent = Inches(0.15)
                paragraph.paragraph_format.first_line_indent = Inches(-0.15)
                paragraph.paragraph_format.space_after = Pt(2)
                add_inline(paragraph, stripped, size=8)
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_inline(paragraph, stripped, size=9.5)
        index += 1

    for section in document.sections:
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].clear()
    footer = document.sections[-1].footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        set_font(run, size=8)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
